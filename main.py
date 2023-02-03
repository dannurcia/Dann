"""Programa principal para la creación de un documento Word con formato GdR
@Dann_Uc - 2022
"""


"""-------------------------------------------------- Módulos -------------------------------------------------------"""

from docx import Document
from datetime import datetime
from docx.enum.section import WD_SECTION
from functions import *


"""-------------------------------------- Variables de configuración de página --------------------------------------"""

top_marg = 2.5
left_marg = 3
right_marg = 3
bottom_marg = 2.5
pag_height = 297
pag_width = 210
head_dist = 1.5
foot_dist = 1


"""------------------------------------- Variables de configuración de párrafos -------------------------------------"""

i = 0                                           # Identacion izquierda del párrafo
fli = 0.5                                       # Identacion primera linea (sangria francesa para valores negativos)
sb = 0                                          # Espaciado antes del párrafo
sa = 6                                          # Espaciado antes del párrafo
ls = 1.2                                        # Interlineado
a = 3                                           # Alineación - 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
fn = 'Calibri(Cuerpo)'                          # Establece el tipo de letra
fs = 12                                         # Establece el tamaño de letra
name_style = 'Indent'                           # Nombre del estilo aplicado


"""----------------------------------- Variables de configuración de 1ra página -------------------------------------"""

nGdR = '001'
dirOfAbrev = 'DIDT'
nameColabAbrev = 'DAUP'
nameColab = 'URCIA PAREDES DANIEL AUGUSTO'
dni = '47856895'
nameEv = 'Implementación de un programa en Python para la etapa de codificación de señales de aceleración'
nameJef = 'Dr. Samuel Gustavo Huamán Bustamante'
date = datetime.today().strftime('%d-%m-%Y')
year = datetime.today().strftime('%Y')
cargo = 'Coordinador del Área de Procesamiento de Señales e IA'
nameInd = 'Construcción de un prototipo de adquisición de señales de aceleración'



"""----------------------------------- Títulos y párrafos del documento ---------------------------------------------"""

titulo1 = 'INTRODUCCIÓN'

parrafo_1 = 'Lorem Ipsum es simplemente el texto de relleno de las imprentas y archivos de texto. Lorem Ipsum ha ' \
            'sido el texto de relleno estándar de las industrias desde el año 1500, cuando un impresor (N. del T. ' \
            'persona que se dedica a la imprenta) desconocido usó una galería de textos y los mezcló de tal manera ' \
            'que logró hacer un libro de textos especimen. No sólo sobrevivió 500 años, sino que tambien ingresó ' \
            'como texto de relleno en documentos electrónicos, quedando esencialmente igual al original. Fue ' \
            'popularizado en los 60s con la creación de las hojas "Letraset", las cuales contenian pasajes de Lorem ' \
            'Ipsum, y más recientemente con software de autoedición, como por ejemplo Aldus PageMaker, el cual ' \
            'incluye versiones de Lorem Ipsum.'

titulo2 = 'ANTECEDENTES'

parrafo_2 = 'Al contrario del pensamiento popular, el texto de Lorem Ipsum no es simplemente texto aleatorio. ' \
            'Tiene sus raices en una pieza cl´sica de la literatura del Latin, que data del año 45 antes de ' \
            'Cristo, haciendo que este adquiera mas de 2000 años de antiguedad. Richard McClintock, un profesor ' \
            'de Latin de la Universidad de Hampden-Sydney en Virginia, encontró una de las palabras más oscuras ' \
            'de la lengua del latín, "consecteur", en un pasaje de Lorem Ipsum, y al seguir leyendo distintos ' \
            'textos del latín, descubrió la fuente indudable. Lorem Ipsum viene de las secciones 1.10.32 y ' \
            '1.10.33 de "de Finnibus Bonorum et Malorum" (Los Extremos del Bien y El Mal) por Cicero, escrito ' \
            'en el año 45 antes de Cristo. Este libro es un tratado de teoría de éticas, muy popular durante el ' \
            'Renacimiento. La primera linea del Lorem Ipsum, "Lorem ipsum dolor sit amet..", viene de una linea ' \
            'en la sección 1.10.32'

titulo3 = 'OBJETIVOS'

parrafo_3 = 'Es un hecho establecido hace demasiado tiempo que un lector se distraerá con el contenido del ' \
            'texto de un sitio mientras que mira su diseño. El punto de usar Lorem Ipsum es que tiene una ' \
            'distribución más o menos normal de las letras, al contrario de usar textos como por ejemplo ' \
            '"Contenido aquí, contenido aquí". Estos textos hacen parecerlo un español que se puede leer. ' \
            'Muchos paquetes de autoedición y editores de páginas web usan el Lorem Ipsum como su texto por ' \
            'defecto, y al hacer una búsqueda de "Lorem Ipsum" va a dar por resultado muchos sitios web que ' \
            'usan este texto si se encuentran en estado de desarrollo. Muchas versiones han evolucionado a ' \
            'través de los años, algunas veces por accidente, otras veces a propósito (por ejemplo insertándole ' \
            'humor y cosas por el estilo).'
titulo31 = 'Objetivo General'
parrafo_31 = 'Es un hecho establecido hace demasiado tiempo que un lector se distraerá con el contenido del ' \
            'texto de un sitio mientras que mira su diseño. El punto de usar Lorem Ipsum es que tiene una ' \
            'distribución más o menos normal de las letras, al contrario de usar textos como por ejemplo ' \
            '"Contenido aquí, contenido aquí". Estos textos hacen parecerlo un español que se puede leer. ' \

titulo32 = 'Objetivos Específicos'
parrafo_32 = 'Es un hecho establecido hace demasiado tiempo que un lector se distraerá con el contenido del ' \
            'texto de un sitio mientras que mira su diseño. El punto de usar Lorem Ipsum es que tiene una ' \
            'distribución más o menos normal de las letras, al contrario de usar textos como por ejemplo ' \
            '"Contenido aquí, contenido aquí". Estos textos hacen parecerlo un español que se puede leer. ' \

objEspecifico321 = 'Este es el primer objetivo especifico.'
objEspecifico322 = 'Este es el segundo objetivo especifico.'
objEspecifico323 = 'Este es el tercer objetivo especifico.'

titulo4 = 'RECURSOS'

parrafo_4 = 'Hay muchas variaciones de los pasajes de Lorem Ipsum disponibles, pero la mayoría sufrió ' \
            'alteraciones en alguna manera, ya sea porque se le agregó humor, o palabras aleatorias que no ' \
            'parecen ni un poco creíbles. Si vas a utilizar un pasaje de Lorem Ipsum, necesitás estar seguro ' \
            'de que no hay nada avergonzante escondido en el medio del texto. Todos los generadores de Lorem ' \
            'Ipsum que se encuentran en Internet tienden a repetir trozos predefinidos cuando sea necesario, ' \
            'haciendo a este el único generador verdadero (válido) en la Internet. Usa un diccionario de mas ' \
            'de 200 palabras provenientes del latín, combinadas con estructuras muy útiles de sentencias, para ' \
            'generar texto de Lorem Ipsum que parezca razonable. Este Lorem Ipsum generado siempre estará libre ' \
            'de repeticiones, humor agregado o palabras no características del lenguaje, etc.'

titulo41 = 'Recursos Humanos'
parrafo_41 = 'Estos son los Recursos Humanos. Estos son los Recursos Humanos. Estos son los Recursos Humanos. ' \
             'Estos son los Recursos Humanos. Estos son los Recursos Humanos.' \

recursohumano411 = 'Este es el primer recurso humano.'
recursohumano412 = 'Este es el segundo recurso humano.'
recursohumano413 = 'Este es el tercer recurso humano.'


titulo42 = 'Recursos materiales'
parrafo_42 = 'Estos son los Recursos Materiales. Estos son los Recursos Materiales. Estos son los ' \
             'Recursos Materiales. Estos son los Recursos Materiales.'
recursomaterial421 = 'Este es el primer recurso material.'
recursomaterial422 = 'Este es el segundo recurso material.'
rec_mat_21 = 'Recurso Material 2.1'
rec_mat_22 = 'Recurso Material 2.2'
rec_mat_23 = 'Recurso Material 2.3'
recursomaterial423 = 'Este es el tercer recurso material.'

titulo43 = 'Otros'
parrafo_43 = 'Estos son otros recursos. Estos son otros recursos. Estos son otros recursos. ' \
             'Estos son otros recursos. Estos son otros recursos. Estos son otros recursos.'

titulo5 = 'ACTIVIDADES DESARROLLADAS'

parrafo_5 = 'Al contrario del pensamiento popular, el texto de Lorem Ipsum no es simplemente texto aleatorio. ' \
            'Tiene sus raices en una pieza cl´sica de la literatura del Latin, que data del año 45 antes de ' \
            'Cristo, haciendo que este adquiera mas de 2000 años de antiguedad. Richard McClintock, un profesor ' \
            'de Latin de la Universidad de Hampden-Sydney en Virginia, encontró una de las palabras más oscuras ' \
            'de la lengua del latín, "consecteur", en un pasaje de Lorem Ipsum, y al seguir leyendo distintos ' \
            'textos del latín, descubrió la fuente indudable. Lorem Ipsum viene de las secciones 1.10.32 y ' \
            '1.10.33 de "de Finnibus Bonorum et Malorum" (Los Extremos del Bien y El Mal) por Cicero, escrito ' \
            'en el año 45 antes de Cristo. Este libro es un tratado de teoría de éticas, muy popular durante el ' \
            'Renacimiento. La primera linea del Lorem Ipsum, "Lorem ipsum dolor sit amet..", viene de una linea ' \
            'en la sección 1.10.32'

titulo51 = 'Actividad 5.1'
parrafo_51 = 'Estas son las Actividades Desarrolladas. Estas son las Actividades Desarrolladas. Estas son las ' \
             'Actividades Desarrolladas. Estas son las Actividades Desarrolladas.'
titulo52 = 'Actividad 5.2'
parrafo_52 = 'Estas son las Actividades Desarrolladas. Estas son las Actividades Desarrolladas. Estas son las ' \
             'Actividades Desarrolladas. Estas son las Actividades Desarrolladas.'
titulo521 = 'Actividad primaria'
parrafo521 = 'Estas son las Actividades Desarrolladas. Estas son las Actividades Desarrolladas. Estas son las ' \
               'Actividades Desarrolladas. Estas son las Actividades Desarrolladas.'
titulo522 = 'Actividad secundaria'
parrafo522 = 'Estas son las Actividades Desarrolladas. Estas son las Actividades Desarrolladas. Estas son las ' \
               'Actividades Desarrolladas. Estas son las Actividades Desarrolladas.'

titulo6 = 'RESULTADOS'

parrafo_6 = 'Al contrario del pensamiento popular, el texto de Lorem Ipsum no es simplemente texto aleatorio. ' \
            'Tiene sus raices en una pieza cl´sica de la literatura del Latin, que data del año 45 antes de ' \
            'Cristo, haciendo que este adquiera mas de 2000 años de antiguedad. Richard McClintock, un profesor ' \
            'de Latin de la Universidad de Hampden-Sydney en Virginia, encontró una de las palabras más oscuras ' \
            'de la lengua del latín, "consecteur", en un pasaje de Lorem Ipsum, y al seguir leyendo distintos ' \
            'textos del latín, descubrió la fuente indudable. Lorem Ipsum viene de las secciones 1.10.32 y ' \
            '1.10.33 de "de Finnibus Bonorum et Malorum" (Los Extremos del Bien y El Mal) por Cicero, escrito ' \
            'en el año 45 antes de Cristo. Este libro es un tratado de teoría de éticas, muy popular durante el ' \
            'Renacimiento. La primera linea del Lorem Ipsum, "Lorem ipsum dolor sit amet..", viene de una linea ' \
            'en la sección 1.10.32'

titulo7 = 'COMENTARIOS'

parrafo_7 = 'Al contrario del pensamiento popular, el texto de Lorem Ipsum no es simplemente texto aleatorio. ' \
            'Tiene sus raices en una pieza cl´sica de la literatura del Latin, que data del año 45 antes de ' \
            'Cristo, haciendo que este adquiera mas de 2000 años de antiguedad. Richard McClintock, un profesor ' \
            'de Latin de la Universidad de Hampden-Sydney en Virginia, encontró una de las palabras más oscuras ' \
            'de la lengua del latín, "consecteur", en un pasaje de Lorem Ipsum, y al seguir leyendo distintos ' \
            'textos del latín, descubrió la fuente indudable. Lorem Ipsum viene de las secciones 1.10.32 y ' \
            '1.10.33 de "de Finnibus Bonorum et Malorum" (Los Extremos del Bien y El Mal) por Cicero, escrito ' \
            'en el año 45 antes de Cristo. Este libro es un tratado de teoría de éticas, muy popular durante el ' \
            'Renacimiento. La primera linea del Lorem Ipsum, "Lorem ipsum dolor sit amet..", viene de una linea ' \
            'en la sección 1.10.32'

titulo8 = 'CONCLUSIONES'

parrafo_8 = 'Estas son las Conclusiones. Estas son las Conclusiones. Estas son las Conclusiones. ' \
            'Estas son las Conclusiones. Estas son las Conclusiones. Estas son las Conclusiones.'

conclu81 = 'Conclusión 1. Conclusión 1. Conclusión 1. Conclusión 1'
conclu82 = 'Conclusión 2. Conclusión 2. Conclusión 2. Conclusión 2.'
conclu83 = 'Conclusión 3. Conclusión 3. Conclusión 3. Conclusión 3.'

titulo9 = 'REFERENCIAS'
ref91 = 'Referencia 1, Referencia 1, Referencia 1, Referencia 1, Referencia 1, Referencia 1, Referencia 1, ' \
        'Referencia 1, Referencia 1, Referencia 1, Referencia 1.'
ref92 = 'Referencia 2, Referencia 2, Referencia 2, Referencia 2, Referencia 2, Referencia 2, Referencia 2, ' \
        'Referencia 2, Referencia 2, Referencia 2, Referencia 2.'
ref93 = 'Referencia 3, Referencia 3, Referencia 3, Referencia 3, Referencia 3, Referencia 3, Referencia 3, ' \
        'Referencia 3, Referencia 3, Referencia 3, Referencia 3.'

titulo10 = 'APÉNDICES'
parrafo_10 = 'Estos son los Apéndices. Estos son los Apéndices. Estos son los Apéndices. Estos son los Apéndices. ' \
             'Estos son los Apéndices. Estos son los Apéndices.'
titulo101 = 'Apéndice 1'
parrafo_101 = 'Estos son los Apéndices. Estos son los Apéndices. Estos son los Apéndices. Estos son los Apéndices. ' \
              'Estos son los Apéndices. Estos son los Apéndices.'
titulo102 = 'Apéndice 2'
parrafo_102 = 'Estos son los Apéndices. Estos son los Apéndices. Estos son los Apéndices. Estos son los Apéndices. ' \
              'Estos son los Apéndices. Estos son los Apéndices.'
titulo11 = 'ANEXOS'
parrafo_11 = 'Estos son los Anexos. Estos son los Anexos. Estos son los Anexos. Estos son los Anexos. ' \
             'Estos son los Anexos. Estos son los Anexos. Estos son los Anexos.'
anexo1 = 'Anexo 1'
anexo2 = 'Anexo 2'
anexo3 = 'Anexo 3'

"""------------------------------------------------- Programa -------------------------------------------------------"""

# Creación del documento
document = Document()

# Primera página
first_page(document, nGdR, dirOfAbrev, nameColabAbrev, nameEv, nameJef, date, cargo, nameInd, year, nameColab, dni)

# Salto de página
document.add_page_break()

# Segunda página
second_page(document)

# Salto de página
document.add_page_break()

# Generación del ÍNDICE
index = document.add_paragraph()
index.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
index.paragraph_format.space_before = Pt(12)
index.paragraph_format.space_after = Pt(24)
run = index.add_run('ÍNDICE')
run.font.size = Pt(18)
run.font.name = 'Arial'
run.bold = True

paragraph = document.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')              # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')       # sets attribute on element
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')      # sets attribute on element
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'    # change 1-3 depending on heading levels you need

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
# fldChar3 = OxmlElement('w:t')
fldChar3 = OxmlElement('w:updateFields')
fldChar3.set(qn('w:val'), 'true')
# fldChar3.text = "Right-click to update field."
fldChar2.append(fldChar3)

fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
p_element = paragraph._p

# Nueva sección e inicio de numeracion de página
document.add_section(WD_SECTION.NEW_PAGE)
document.sections[1].footer.is_linked_to_previous = False
add_page_number(document.sections[1])

# Llamada a la función de estilo de párrafo
paragraph_0(document, i, fli, sb, sa, ls, a, fn, fs, name_style)

# Creación del texto del documento (títulos y párrafos)
title_0(document, titulo1, 1)
p1 = document.add_paragraph(parrafo_1)

title_0(document, titulo2, 2)
p2 = document.add_paragraph(parrafo_2)

title_0(document, titulo3, 3)
p3 = document.add_paragraph(parrafo_3)
title_1(document, titulo31, 3.1)
p31 = document.add_paragraph(parrafo_31)
title_1(document, titulo32, 3.2)
# p32 = document.add_paragraph(parrafo_32)
item321 = items(document, objEspecifico321)
item322 = items(document, objEspecifico322)
item323 = items(document, objEspecifico323)

title_0(document, titulo4, 4)

p4 = document.add_paragraph(parrafo_4)
title_1(document, titulo41, 4.1)
p41 = document.add_paragraph(parrafo_41)
rh411 = items(document, recursohumano411)
rh412 = items(document, recursohumano412)
rh413 = items(document, recursohumano413)
title_1(document, titulo42, 4.2)
p42 = document.add_paragraph(parrafo_42)
rm421 = items(document, recursomaterial421)
rm422 = items(document, recursomaterial422)
rm4221 = bullets(document, rec_mat_21)
rm4222 = bullets(document, rec_mat_22)
rm4223 = bullets(document, rec_mat_23)
rm423 = items(document, recursomaterial423)
title_1(document, titulo43, 4.3)
p43 = document.add_paragraph(parrafo_43)

title_0(document, titulo5, 5)
p5 = document.add_paragraph(parrafo_5)
title_1(document, titulo51, 5.1)
p51 = document.add_paragraph(parrafo_51)
contenido_tabla = "Cabecera 1$$Cabecera 2$$Cabecera 3 es mas grande$$Cabecera 4||C11$$C21$$C31$$C41||C12$$C22$$C32$$C42||C13$$C23$$C33$$C43"
get_tabla(document, 1, "Titulo de la Tabla", contenido_tabla)
title_1(document, titulo52, 5.2)
p52 = document.add_paragraph(parrafo_52)
title_2(document, titulo521, 5.2, 1)
p521 = document.add_paragraph(parrafo521)
get_imagen(document, r"C:\Users\user\Desktop\Dann Uc\Sistema de gestión de documentos\Pruebas\roles.png", 9.54, 1, "Titulo de la figura")
title_2(document, titulo522, 5.2, 2)
p522 = document.add_paragraph(parrafo522)

title_0(document, titulo6, 6)
p6 = document.add_paragraph(parrafo_6)

title_0(document, titulo7, 7)
p7 = document.add_paragraph(parrafo_7)

title_0(document, titulo8, 8)
p8 = document.add_paragraph(parrafo_8)
c1 = items(document, conclu81)
c2 = items(document, conclu82)
c3 = items(document, conclu83)

title_0(document, titulo9, 9)
ref1 = ref_bullets(document, ref91, 1)
ref2 = ref_bullets(document, ref92, 2)
ref3 = ref_bullets(document, ref92, 3)

title_0(document, titulo10, 10)
p10 = document.add_paragraph(parrafo_10)
title_1(document, titulo101, 10.1)
p101 = document.add_paragraph(parrafo_101)
title_1(document, titulo102, 10.2)
p102 = document.add_paragraph(parrafo_102)

title_0(document, titulo11, 11)
p11 = document.add_paragraph(parrafo_11)
anex1 = items(document, anexo1)
anex2 = items(document, anexo2)
anex3 = items(document, anexo3)

# Aplicación del estilo al párrafo creado
p1.style = document.styles[name_style]
p2.style = document.styles[name_style]
p3.style = document.styles[name_style]
p31.style = document.styles[name_style]
# p32.style = document.styles[name_style]
p4.style = document.styles[name_style]
p41.style = document.styles[name_style]
p42.style = document.styles[name_style]
p43.style = document.styles[name_style]
p5.style = document.styles[name_style]
p51.style = document.styles[name_style]
p52.style = document.styles[name_style]
p521.style = document.styles[name_style]
p522.style = document.styles[name_style]
p6.style = document.styles[name_style]
p7.style = document.styles[name_style]
p8.style = document.styles[name_style]
p10.style = document.styles[name_style]
p101.style = document.styles[name_style]
p102.style = document.styles[name_style]
p11.style = document.styles[name_style]

# Llamada a la función de formato de página
page_format(document, top_marg, left_marg, right_marg, bottom_marg, pag_height, pag_width, head_dist, foot_dist)

# Guardado del documento
document.save(r'C:\Users\user\Desktop\Dann Uc\Sistema de gestión de documentos\Pruebas\Sistema_Gestion_Documentario.docx')
