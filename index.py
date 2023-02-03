
from docx import Document
from docx.shared import Inches, Pt, Cm, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from functions import *

def del_paragraph(paragraph):
    """

    :param paragraph:
    :return:
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def indice(ret, word):
    """
    Formato de las palabras del indice
    :return:
    """

    if ret == 'title_0':
        i = document.add_paragraph()
        run = i.add_run(word)
        i.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        i.paragraph_format.line_spacing = 1.4
        i.paragraph_format.left_indent = Cm(0)
        i.paragraph_format.space_before = Pt(6)
        i.paragraph_format.space_after = Pt(6)
        # i.paragraph_format.first_line_indent = Cm(0.63)
        run.font.name = 'Calibri Light'
        run.font.size = Pt(12)
        run.bold = True

    elif ret == 'title_1':
        i = document.add_paragraph()
        run = i.add_run(word)
        i.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        i.paragraph_format.line_spacing = 1.4
        i.paragraph_format.left_indent = Cm(0.5)
        i.paragraph_format.space_before = Pt(6)
        i.paragraph_format.space_after = Pt(6)
        # i.paragraph_format.first_line_indent = Cm(0.63)
        run.font.name = 'Calibri(Cuerpo)'
        run.font.size = Pt(12)
        run.bold = False


document = Document()

titulo1 = 'INTRODUCCIÓN'
titulo2 = 'ANTECEDENTES'
titulo3 = 'OBJETIVOS'
titulo31 = 'Objetivo General'
titulo32 = 'Objetivos Específicos'
titulo4 = 'RECURSOS'
titulo41 = 'Recursos Humanos'
titulo42 = 'Recursos materiales'
titulo43 = 'Otros'
titulo5 = 'ACTIVIDADES DESARROLLADAS'
titulo51 = 'Actividad 1'
titulo52 = 'Actividad 2'
titulo6 = 'RESULTADOS'
titulo7 = 'COMENTARIOS'
titulo8 = 'CONCLUSIONES'
titulo9 = 'REFERENCIAS'
titulo10 = 'APÉNDICES'
titulo101 = 'Apéndice 1'
titulo102 = 'Apéndice 2'
titulo11 = 'ANEXOS'

# Obtención de las variables de retorno
ret1, word1 = title_0(document, titulo1, 1)
print(ret1, word1)
ret2, word2 = title_0(document, titulo2, 2)
ret3, word3 = title_0(document, titulo3, 3)
ret31, word31 = title_1(document, titulo31, 3.1)
print(ret31, word31)
ret32, word32 = title_1(document, titulo32, 3.2)
ret4, word4 = title_0(document, titulo4, 4)
ret41, word41 = title_1(document, titulo41, 4.1)
ret42, word42 = title_1(document, titulo42, 4.2)
ret43, word43 = title_1(document, titulo43, 4.3)
ret5, word5 = title_0(document, titulo5, 5)
ret51, word51 = title_1(document, titulo51, 5.1)
ret52, word52 = title_1(document, titulo52, 5.2)
ret6, word6 = title_0(document, titulo6, 6)
ret7, word7 = title_0(document, titulo7, 7)
ret8, word8 = title_0(document, titulo8, 8)
ret9, word9 = title_0(document, titulo9, 9)
ret10, word10 = title_0(document, titulo10, 10)
ret101, word101 = title_1(document, titulo101, 10.1)
ret102, word102 = title_1(document, titulo102, 10.2)
ret11, word11 = title_0(document, titulo11, 11)

# # Generación del indice
# index = document.add_paragraph()
# index.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# index.paragraph_format.space_before = Pt(12)
# index.paragraph_format.space_after = Pt(24)
# run = index.add_run('ÍNDICE')
# run.font.size = Pt(18)
# run.font.name = 'Arial'
# run.bold = True
#
# indice(ret1, word1)
# indice(ret2, word2)
# indice(ret3, word3)
# indice(ret31, word31)
# indice(ret32, word32)
# indice(ret4, word4)
# indice(ret41, word41)
# indice(ret42, word42)
# indice(ret43, word43)
# indice(ret5, word5)
# indice(ret51, word51)
# indice(ret52, word52)
# indice(ret6, word6)
# indice(ret7, word7)
# indice(ret8, word8)
# indice(ret9, word9)
# indice(ret10, word10)
# indice(ret101, word101)
# indice(ret102, word102)
# indice(ret11, word11)

# Guardado del documento
document.save(r'C:\Users\user\Desktop\Dann Uc\Sistema de gestión de documentos\Pruebas\index.docx')