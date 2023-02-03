from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

document = Document()


def get_tabla(docum, num_tabla, tit_tabla, lista_elem):
    # PONIENDO EL TITULO
    parrafo = docum.add_paragraph('Tabla ' + str(num_tabla) + ". " + tit_tabla)
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.paragraph_format.space_before = Pt(12)
    parrafo.paragraph_format.space_after = Pt(6)
    ru = parrafo.runs[0]
    ru.font.size = Pt(12)
    ru.font.name = 'Calibri'

    vvv = lista_elem.split("||")
    # definiendo la dimension de la tabla
    num_rows = len(vvv)
    num_cols = len(vvv[0].split("$$"))
    tablita = docum.add_table(num_rows, num_cols)
    tablita.autofit = True
    tablita.alignment = WD_TABLE_ALIGNMENT.CENTER
    cont = 0
    for k in vvv:
        fff = k.split("$$")
        rrr = tablita.rows[cont].cells
        n_o = 0
        for h in fff:

            rrr[n_o].text = h

            ppp = rrr[n_o].paragraphs[0]
            ppp.paragraph_format.space_before = Pt(0)
            ppp.paragraph_format.space_after = Pt(6)
            mmm = ppp.runs[0]
            mmm.font.size = Pt(12)
            mmm.font.name = 'Calibri'

            if (cont == 0):
                ppp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            n_o = n_o + 1

        cont = cont + 1

    tablita.style = 'Table Grid'

    # agregando un espaciado despues de la tabla
    p_esp = docum.add_paragraph()
    p_esp.paragraph_format.space_before = Pt(12)
    p_esp.paragraph_format.space_after = Pt(0)
    r_esp = p_esp.add_run('')
    r_esp.font.size = Pt(12)
    r_esp.font.name = 'Calibri'

    docum.save(r"C:\Users\user\Desktop\Dann Uc\Sistema de gestión de documentos\Pruebas\neworder.docx")


contenido = "Cabecera 1$$Cabecera 2$$Cabecera 3 es mas grande$$Cabecera 4||C11$$C21$$C31$$C41||C12$$C22$$C32$$C42||C13$$C23$$C33$$C43"
get_tabla(document, 1, "Titulo de la Tabla", contenido)
