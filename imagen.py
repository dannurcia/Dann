from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
import os
from docx import Document

# path_b = os.path.dirname(os.path.abspath(__file__))
document = Document()


def get_imagen(docum, path_img, ancho_im, num_im, tit_im):
    #agregando un espaciado antes de la imagen
    p_esp = docum.add_paragraph()
    p_esp.paragraph_format.space_before = Pt(12)
    p_esp.paragraph_format.space_after = Pt(0)
    r_esp = p_esp.add_run('')
    r_esp.font.size = Pt(12)
    r_esp.font.name = 'Calibri'

    docum.add_picture(path_img, width = Cm(ancho_im))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Poniendo el titulo
    parrafo = docum.add_paragraph()
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.paragraph_format.space_before = Pt(0)
    parrafo.paragraph_format.space_after = Pt(6)
    r = parrafo.add_run('Figura ' + str(num_im) + ". " + tit_im)
    r.font.size = Pt(12)
    r.font.name = 'Calibri'

get_imagen(document, r"C:\Users\user\Desktop\Dann Uc\Sistema de gestión de documentos\Pruebas\roles.png", 9.54, 1, "Titulo de la figura")

document.save(r"C:\Users\user\Desktop\Dann Uc\Sistema de gestión de documentos\Pruebas\wordim.docx")