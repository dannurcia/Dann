from database.pg_connection_database import DatabaseConnection
from docx import Document
# from docx.shared import Cm
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.shared import Inches
# from docx.shared import Pt
# from datetime import datetime
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# import sys
from docx.enum.section import WD_SECTION
from functions import *
import re
from docx2pdf import convert

"""----------------------------------- Variables de configuración de 1ra página -------------------------------------"""

dirOfAbrev = 'DIDT'
nameEv = 'Implementación de un programa en Python para la etapa de codificación de señales de aceleración'
nameJef = 'Dr. Samuel Gustavo Huamán Bustamante'
date = datetime.today().strftime('%d-%m-%Y')
year = datetime.today().strftime('%Y')
mes = datetime.today().strftime('%m')
mes_actual = mes2words(mes)
cargo = 'Coordinador del Área de Procesamiento de Señales e IA'


"""-------------------------------------- Variables de configuración de página --------------------------------------"""

top_marg = 2.5
left_marg = 3
right_marg = 3
bottom_marg = 1.5
pag_height = 297
pag_width = 210
head_dist = 1.5
foot_dist = 1


"""------------------------------------- Variables de configuración de párrafos -------------------------------------"""

i = 0                                           # Identacion izquierda del párrafo
fli = 0.5                                       # Identacion primera linea (sangria francesa para valores negativos)
sb = 0                                          # Espaciado antes del párrafo
sa = 6                                          # Espaciado antes del párrafo
ls = 1.2                                        # Interlineado
a = 3                                           # Alineación - 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
fn = 'Calibri'                                  # Establece el tipo de letra
fs = 12                                         # Establece el tamaño de letra
name_style = 'Indent'                           # Nombre del estilo aplicado



def create_document(id_user, id_document, id_evidencia):
    # -------------------------------------------------------------
    # Obtener el nombre completo de usuario, sus iniciales y su DNI
    # -------------------------------------------------------------
    records, msg = db_link.get_user_info(id_user)
    print(msg)
    if records:
        full_name = records[0][1] + ' ' + records[0][2]
        name_initial = records[0][1].split() + records[0][2].split()
        name_initials = ''.join([n[0] for n in name_initial])
        dni = records[0][3]
        # print(full_name)
        # print(name_initials)
        # print(dni)
    else:
        print(msg)
    # -------------------------------------------------------------
    # Obtener el texto del documento a partir del formulario
    # -------------------------------------------------------------
    body_text = []
    # Introducción
    body_text.append('INTRODUCCIÓN')
    records, msg = db_link.get_introduccion(id_document)
    if records:
        text_prueba = records[0][0]
        # text_prueba = re.sub(r'\r\r', '', text_prueba)
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
        print(body_text)
        # print(text_prueba)
    else:
        print(msg)

    # Antecedentes
    body_text.append('ANTECEDENTES')
    records, msg = db_link.get_antecedentes(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Objetivos
    body_text.append('OBJETIVOS')
    # records, msg = db_link.get_objectivos(id_document)
    # if records:
    #     body_text.append(records[0][0])
    # else:
    #     print(msg)

    # Objetivo General
    body_text.append('Objetivo General')
    records, msg = db_link.get_objectivo_general(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Objetivos Específicos
    body_text.append('Objetivos Específicos')
    records, msg = db_link.get_objectivos_especificos(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Recursos
    body_text.append('RECURSOS')
    # records, msg = db_link.get_recursos(id_document)
    # if records:
    #     body_text.append(records[0][0])
    # else:
    #     print(msg)

    # Recursos Humanos
    body_text.append('Recursos Humanos')
    records, msg = db_link.get_recursos_humanos(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Recursos Materiales
    body_text.append('Recursos Materiales')
    records, msg = db_link.get_recursos_materiales(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Otros
    body_text.append('Otros')
    records, msg = db_link.get_otros_recursos(id_document)
    if records:
        body_text.append(records[0][0])
    else:
        print(msg)

    # Actividades Desarrolladas
    body_text.append('ACTIVIDADES DESARROLLADAS')
    records, msg = db_link.get_cuerpo_evidencia(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Resultados
    body_text.append('RESULTADOS')
    records, msg = db_link.get_resultados(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Comentarios
    body_text.append('COMENTARIOS')
    records, msg = db_link.get_comentarios(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Conclusiones
    body_text.append('CONCLUSIONES')
    records, msg = db_link.get_conclusiones(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Referencias
    body_text.append('REFERENCIAS')
    records, msg = db_link.get_referencias(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Apéndices
    body_text.append('APÉNDICES')
    records, msg = db_link.get_apendices(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    # Anexos
    body_text.append('ANEXOS')
    records, msg = db_link.get_anexos(id_document)
    if records:
        text_prueba = records[0][0]
        text_prueba_list = text_prueba.split('\r\n')
        body_text.append(text_prueba_list)
    else:
        print(msg)

    """ Corregir nEvi dado que ya no se usara para generar el numero de informe sino de una tabla de fechas"""

    # Numero de evidencia
    nEvi = str(id_evidencia)
    # print(nEvi)
    if len(nEvi) == 1:
        nEvi_form = '00' + nEvi
    elif len(nEvi) == 2:
        nEvi_form = '0' + nEvi
    else:
        pass

    # # records, msg = db_link.get_cod_evidencia(id_evidencia)
    # if records:
    #     nEvi = str(id_evidencia)
    #     # print(nEvi)
    #     if len(nEvi) == 1:
    #         nEvi_form = '00' + nEvi
    #     elif len(nEvi) == 2:
    #         nEvi_form = '0' + nEvi
    #     else:
    #         pass
    # else:
    #     pr

    # Tabla evidenciaGdR
    descripcion_evi, msg = db_link.get_descripcion_eviGdR(nEvi)
    descripcion_evi = descripcion_evi[0][0]
    if descripcion_evi:
        pass
        # print('Tabla evidenciaGdR: ', descripcion_evi)
    else:
        print(msg)

    # Tabla prioridadGdR
    records_pri, msg = db_link.get_prioridadGdR(nEvi)
    # print('records_pri:', records_pri)
    if records:
        pass
        # print('Tabla prioridadGdR: ', records_pri)
    else:
        print(msg)

    # Descripcion del indicador
    records_ind, msg = db_link.get_descripcion_indGdR(nEvi)
    records_ind = records_ind[0][0]
    print('records_ind:', records_ind)
    if records:
        pass
        # print('Tabla prioridadGdR: ', records_pri)
    else:
        print(msg)



    # print(*body_text, sep='\n')

    # -------------------------------------------------------------
    # Crear documento word
    # -------------------------------------------------------------

    # Creando el documento
    document = Document()

    first_page(document, nEvi_form, dirOfAbrev, name_initials, nameEv, nameJef, date, cargo, year, dni, records_ind, descripcion_evi, full_name)

    # Salto de página
    document.add_page_break()

    # Segunda página
    second_page(document, descripcion_evi, full_name, year, mes_actual)

    # Salto de página
    document.add_page_break()

    # Generación del ÍNDICE
    index = document.add_paragraph()
    index.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    index.paragraph_format.space_before = Pt(12)
    index.paragraph_format.space_after = Pt(24)
    run = index.add_run('ÍNDICE')
    run.font.size = Pt(18)
    run.font.name = 'Arial'
    run.bold = True

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    # fldChar3 = OxmlElement('w:t')
    fldChar3 = OxmlElement('w:updateFields')
    fldChar3.set(qn('w:val'), 'true')
    # fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

    # Nueva sección e inicio de numeracion de página
    document.add_section(WD_SECTION.NEW_PAGE)
    document.sections[1].footer.is_linked_to_previous = False
    add_page_number(document.sections[1])


    ############## AÑADIDO: ######################

    # Llamada a la función de estilo de párrafo
    paragraph_0(document, i, fli, sb, sa, ls, a, fn, fs, name_style)

    # Creacion del texto del documento (titulos y parrafos)
    title_0(document, body_text[0], 1) # Introduccion
    print(body_text)
    print(len(body_text))
    for parr in body_text[1]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p1 = document.add_paragraph(body_text[1])

    title_0(document, body_text[2], 2) # Antecedentes
    for parr in body_text[3]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p2 = document.add_paragraph(body_text[3])

    title_0(document, body_text[4], 3) # Objetivos
    # p3 = document.add_paragraph(body_text[3])
    title_1(document, body_text[5], 3.1) # Obj. general
    for parr in body_text[6]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p31 = document.add_paragraph(body_text[6])
    title_1(document, body_text[7], 3.2) # Obj. especifico
    for parr in body_text[8]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p311 = document.add_paragraph(body_text[8])
    # oe311 = items(document, body_text[8])
    # oe312 = items(document, body_text[8])
    # oe313 = items(document, body_text[8])

    title_0(document, body_text[9], 4) # Recursos
    title_1(document, body_text[10], 4.1) # Recursos humanos
    for parr in body_text[11]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p41 = document.add_paragraph(body_text[11])
    title_1(document, body_text[12], 4.2) # Recursos materiales
    for parr in body_text[13]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p42 = document.add_paragraph(body_text[13])
    title_1(document, body_text[14], 4.3) # Otros
    for parr in body_text[15]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p42 = document.add_paragraph(body_text[13])

    title_0(document, body_text[16], 5)  # Actividades desarrolladas
    for parr in body_text[17]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p5 = document.add_paragraph(body_text[17])

    title_0(document, body_text[18], 6)  # Resultados
    for parr in body_text[19]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p6 = document.add_paragraph(body_text[19])

    title_0(document, body_text[20], 7)  # Comentarios
    for parr in body_text[21]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p7 = document.add_paragraph(body_text[21])

    title_0(document, body_text[22], 8)  # Conclusiones
    for parr in body_text[23]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p8 = document.add_paragraph(body_text[23])

    title_0(document, body_text[24], 9)  # Referencias
    for parr in body_text[25]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p9 = document.add_paragraph(body_text[25])

    title_0(document, body_text[26], 10)  # Apendices
    for parr in body_text[27]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p10 = document.add_paragraph(body_text[27])

    title_0(document, body_text[28], 11)  # Anexos
    for parr in body_text[29]:
        parrs = document.add_paragraph(parr)
        parrs.style = document.styles[name_style]
    # p11 = document.add_paragraph(body_text[29])

    # Llamada a la función de formato de página
    page_format(document, top_marg, left_marg, right_marg, bottom_marg, pag_height, pag_width, head_dist, foot_dist)

    ############# FIN AÑADIDO #####################


    document_path = 'C:\\Users\\user\\Desktop\\Dann Uc\\Sistema de gestión de documentos\\Pruebas\\documento' + str(id_document[0][0][0]) + '.docx'
    document.save(document_path)
    document_path = 'documento' + str(id_document[0][0][0]) + '.docx'

    # convert('C:\\Users\\user\\Desktop\\Dann Uc\\Sistema de gestión de documentos\\Pruebas\\documento' + str(id_document) + '.docx', 'C:\\Users\\user\\Desktop\\Dann Uc\\Sistema de gestión de documentos\\Pruebas\\documento' + str(id_document) + '.pdf')

    return document_path


if __name__ == '__main__':
    db_link = DatabaseConnection()
    # id_user = int(sys.argv[1])
    # id_document = int(sys.argv[2])
    id_user = 12
    id_evidencia = 1
    id_documento = db_link.get_documento(id_evidencia)
    document_path = create_document(id_user, id_documento, id_evidencia)
    print(document_path, end='')
    db_link.close_connection()

