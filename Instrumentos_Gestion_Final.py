"""Programa principal para la creación del documento Word Resumen Ejecutivo de un Proyecto
@Dann_Uc - 2022
"""


"""-------------------------------------------------- Módulos -------------------------------------------------------"""

from docx import Document
from docx.enum.section import WD_SECTION
from functions import *
from docx.enum.text import WD_ALIGN_PARAGRAPH


"""-------------------------------------- Variables de configuración de página --------------------------------------"""

top_marg = 2.5
left_marg = 3
right_marg = 3
bottom_marg = 2.5
pag_height = 297
pag_width = 210
head_dist = 1.5
foot_dist = 1


"""------------------------------------- Variables de configuración de párrafos -------------------------------------"""

i = 0                                           # Identacion izquierda del párrafo
fli = 0.5                                       # Identacion primera linea (sangria francesa para valores negativos)
sb = 0                                          # Espaciado antes del párrafo
sa = 6                                          # Espaciado antes del párrafo
ls = 1.2                                        # Interlineado
a = 3                                           # Alineación - 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
fn = 'Arial'                                    # Establece el tipo de letra
fs = 10                                          # Establece el tamaño de letra
name_style = 'Indent'                           # Nombre del estilo aplicado


"""-------------------------------------------- Variables del documento ---------------------------------------------"""

titles = ['Datos Generales del Proyecto', 'Entidades Colaboradoras', 'Equipo del Proyecto',
          'Información del Proyecto de Investigación', 'Información del Proyecto de Desarrollo Tecnológico',
          'Información del Proyecto de Innovación', 'Memoria Técnica', 'Presupuesto']

subtitles = {titles[0]: ['Título del proyecto', 'Resumen ejecutivo', 'Duración del proyecto', 'Área y línea de investigación',
                         'Localización del proyecto', 'Datos de los investigadores (CV descargado del CTI Vitae)',
                         'Tipo de proyecto (Investigación o Desarrollo)'],
             titles[1]: ['Nombre de entidad', 'Tipo de entidad', 'RUC', 'Teléfono', 'Correo'],
             titles[2]: ['Coordinador(a) del proyecto', 'Coordinador(a) administrativa', 'Investigadores(as)',
                         'Desarrolladores(as)', 'Técnicos'],
             titles[3]: ['Referencia de la investigación en la que se fundamenta',
                         'Conocimiento que se desea obtener o área que fortalecer'],
             titles[4]: ['Referencia de la investigación en la que se fundamenta',
                         'Producto o prototipo que se desea desarrollar para transferencia',
                         'Mercado o empresa objetivo'],
             titles[5]: ['Planteamiento del Problema de Investigación',
                         'Situación Problemática', 'Formulación del Problema',
                         'Justificación',
                         'Objetivos (General y especifico)',
                         'Limitaciones', 'Marco Teórico', 'Antecedentes del Problema', 'Bases teóricas o marco conceptual',
                         'Hipótesis y variables', 'Formulación de hipótesis', 'Variables', 'Metodología del proyecto',
                         'Diseño metodológico', 'Técnicas estadísticas para el procesamiento de datos',
                         'Describa los riesgos del proyecto', 'Aspectos éticos y regulatorios',
                         'Impactos esperados'],
             titles[6]: ['Cronograma de actividades', 'Presupuesto', 'Fuentes de Financiamiento'],
             titles[7]: ['Resultados Esperados', 'Referencia Bibliografías', 'Anexos']}

parrafo_1 = 'Guia metodologica para la formulación de proyectos de investigación y desarrollo tecnológico con ' \
           'criterios establecidos por el ente rector en ciencia y tecnología.'

parrafo_2 = 'Para la realización de la presente guia metodologica se ha recibido y tomado en consideración los ' \
            'aportes de los Coordinadores de la Dirección de Investigación y Desarrollo Tecnologico del INICTEL-UNI. ' \
            'Uno de los documentos empleado como referencia para la propuesta fue la Guía Práctica para la Formulación ' \
            'y Ejecución de Proyectos de Investigación y Desarrollo (I+D) propuesta por el CONCYTEC , la misma que ' \
            'busca establecer pautas mínimas para estandarizar la estructura, evaluación y gestión  de proyectos de ' \
            'I+D de las instituciones que como el INICTEL-UNI componen el Sistema Nacional de Ciencia Tecnología e ' \
            'Innovación Tecnológica (SINACYT) [1].'

"""------------------------------------------------- Programa -------------------------------------------------------"""

# Creación del documento
document = Document()

# Introduciendo el logo del encabezado
header = document.sections[0].header
paragraph = header.paragraphs[0]
logo_run = paragraph.add_run()
logo_run.add_picture("cabecera.png", width=Cm(15))

# Inicio de numeración de página
add_page_number(document.sections[0])

# Llamada a la función que crea la 1ra página
first_page_re(document)

# Salto de página
document.add_page_break()

# Llamada a la función de estilo de párrafo
paragraph_re(document, i, fli, sb, sa, ls, a, fn, fs, name_style)

# Título del cuerpo del documento
heading = document.add_heading('GUIA METODOLOGICA PARA LA FORMULACIÓN DE PROYECTOS DE INVESTIGACIÓN Y DESARROLLO '
                               'TECNOLOGICO EN DIDT', level=2)
title_style = heading.style
rFonts = title_style.element.rPr.rFonts  # lxml de la fuente
rFonts.set(qn("w:asciiTheme"), "Arial")  # Configura fuente xml
title_style.font.color.rgb = RGBColor(0x00, 0x70, 0xc0)
title_style.paragraph_format.left_indent = Cm(0.63)
title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_style.paragraph_format.space_before = Pt(6)
title_style.paragraph_format.space_after = Pt(6)
font = title_style.font  # Estilo de fuente del titulo
font.name = 'Arial'  # Necesario, acorde al xml
font.size = Pt(14)  # Tamaño de fuente
font.bold = True

r = document.add_paragraph()
paragraph_format = r.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = r.add_run('\nINDICADOR / PRODUCTO')
r1.bold, r1.underline = True, False
r1.font.size, r1.font.name = Pt(10), 'Arial'

s = document.add_paragraph(parrafo_1)
s.style = document.styles[name_style]

t = document.add_paragraph()
paragraph_format = t.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
t1 = t.add_run('\nGUIA METODOLOGICA PARA LA FORMULACIÓN DE PROYECTOS DE INVESTIGACIÓN Y DESARROLLO TECNOLÓGICO '
               'CON CRITERIOS ESTABLECIDOS POR EL ENTE RECTOR EN CIENCIA Y TECNOLOGÍA')
t1.bold, t1.underline = True, False
t1.font.size, t1.font.name = Pt(10), 'Arial'

u = document.add_paragraph(parrafo_2)
u.style = document.styles[name_style]

titles_0(document, 'Propuesta de Formato Estándar de Proyecto de Investigación', 1)
titles_1(document, 'Datos de Identificación del Proyecto', 'I')
titles_2(document, subtitles[titles[0]][0], 1.1)
titles_2(document, subtitles[titles[0]][1], 1.2)
titles_2(document, subtitles[titles[0]][2], 1.3)
titles_2(document, subtitles[titles[0]][3], 1.4)
titles_2(document, subtitles[titles[0]][4], 1.5)
titles_2(document, subtitles[titles[0]][5], 1.6)
titles_2(document, subtitles[titles[0]][6], 1.7)
titles_1(document, 'Entidades Colaboradoras', 'II')
titles_2(document, subtitles[titles[1]][0], 2.1)
titles_2(document, subtitles[titles[1]][1], 2.2)
titles_2(document, subtitles[titles[1]][2], 2.3)
titles_2(document, subtitles[titles[1]][3], 2.4)
titles_2(document, subtitles[titles[1]][4], 2.5)
titles_1(document, 'Equipo del Proyecto', 'III')
titles_2(document, subtitles[titles[2]][0], 3.1)
titles_2(document, subtitles[titles[2]][1], 3.2)
titles_2(document, subtitles[titles[2]][2], 3.3)
titles_2(document, subtitles[titles[2]][3], 3.4)
titles_2(document, subtitles[titles[2]][4], 3.5)
titles_1(document, 'Información del proyecto de investigación (Si fuera el caso)', 'IV')
titles_2(document, subtitles[titles[3]][0], 4.1)
titles_2(document, subtitles[titles[3]][1], 4.2)
titles_1(document, 'Información del proyecto de desarrollo tecnológico (Si fuera el caso)', 'V')
titles_2(document, subtitles[titles[4]][0], 5.1)
titles_2(document, subtitles[titles[4]][1], 5.2)
titles_2(document, subtitles[titles[4]][2], 5.3)
titles_1(document, 'VI.	Descripción del Proyecto', 'VI')
titles_2(document, subtitles[titles[5]][0], 6.1)
titles_3(document, subtitles[titles[5]][1], 6.1, 1)
titles_3(document, subtitles[titles[5]][2], 6.1, 2)
titles_3(document, subtitles[titles[5]][3], 6.1, 3)
titles_3(document, subtitles[titles[5]][4], 6.1, 4)
titles_3(document, subtitles[titles[5]][5], 6.1, 5)
titles_2(document, subtitles[titles[5]][6], 6.2)
titles_3(document, subtitles[titles[5]][7], 6.2, 1)
titles_3(document, subtitles[titles[5]][8], 6.2, 2)
titles_2(document, subtitles[titles[5]][9], 6.3)
titles_3(document, subtitles[titles[5]][10], 6.3, 1)
titles_3(document, subtitles[titles[5]][11], 6.3, 2)
titles_2(document, subtitles[titles[5]][12], 6.4)
titles_3(document, subtitles[titles[5]][13], 6.4, 1)
titles_3(document, subtitles[titles[5]][14], 6.4, 2)
titles_3(document, subtitles[titles[5]][15], 6.4, 3)
titles_3(document, subtitles[titles[5]][16], 6.4, 4)
titles_3(document, subtitles[titles[5]][17], 6.4, 5)










# Cuerpo del documento (títulos, subtítulos y párrafos)
# p = document.add_paragraph()
# p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# p.paragraph_format.space_before = Pt(12)                                      # Espaciado 12 pts antes del párrafo
# p.paragraph_format.space_after = Pt(20)
# p.paragraph_format.line_spacing = 1.5
# p1 = p.add_run('Formato estándar que permita definir el tipo de proyecto')
# p1.bold, p1.font.size, p1.font.name = True, Pt(10), 'Arial'

# titles_0(document, titles[0], 1)
# titles_1(document, subtitles[titles[0]][0], 1.1)
# titles_1(document, subtitles[titles[0]][1], 1.2)
# titles_1(document, subtitles[titles[0]][2], 1.3)
# titles_1(document, subtitles[titles[0]][3], 1.4)
# titles_1(document, subtitles[titles[0]][4], 1.5)
#
# titles_0(document, titles[1], 2)
# titles_1(document, subtitles[titles[1]][0], 2.1)
# titles_1(document, subtitles[titles[1]][1], 2.2)
# titles_1(document, subtitles[titles[1]][2], 2.3)
# titles_1(document, subtitles[titles[1]][3], 2.4)
# titles_1(document, subtitles[titles[1]][4], 2.5)
#
# titles_0(document, titles[2], 3)
# titles_1(document, subtitles[titles[2]][0], 3.1)
# titles_1(document, subtitles[titles[2]][1], 3.2)
# titles_1(document, subtitles[titles[2]][2], 3.3)
# titles_1(document, subtitles[titles[2]][3], 3.4)
# titles_1(document, subtitles[titles[2]][4], 3.5)
#
# titles_0(document, titles[3], 4)
# titles_1(document, subtitles[titles[3]][0], 4.1)
# titles_1(document, subtitles[titles[3]][1], 4.2)
#
# titles_0(document, titles[4], 5)
# titles_1(document, subtitles[titles[4]][0], 5.1)
# titles_1(document, subtitles[titles[4]][1], 5.2)
# titles_1(document, subtitles[titles[4]][2], 5.3)
#
# titles_0(document, titles[5], 6)
# titles_1(document, subtitles[titles[5]][0], 6.1)
# titles_1(document, subtitles[titles[5]][1], 6.2)
# titles_1(document, subtitles[titles[5]][2], 6.3)
# titles_1(document, subtitles[titles[5]][3], 6.4)
# titles_1(document, subtitles[titles[5]][4], 6.5)
# titles_1(document, subtitles[titles[5]][5], 6.6)
#
# titles_0(document, titles[6], 7)
# titles_1(document, subtitles[titles[6]][0], 7.1)
# titles_1(document, subtitles[titles[6]][1], 7.2)
# titles_1(document, subtitles[titles[6]][2], 7.3)
# titles_1(document, subtitles[titles[6]][3], 7.4)
# titles_1(document, subtitles[titles[6]][4], 7.5)
# titles_1(document, subtitles[titles[6]][5], 7.6)
#
# titles_0(document, titles[7], 8)
# titles_1(document, subtitles[titles[7]][0], 8.1)
# titles_1(document, subtitles[titles[7]][1], 8.2)
# titles_1(document, subtitles[titles[7]][2], 8.3)
# titles_1(document, subtitles[titles[7]][3], 8.4)
# titles_1(document, subtitles[titles[7]][4], 8.5)
# titles_1(document, subtitles[titles[7]][5], 8.6)
# titles_1(document, subtitles[titles[7]][6], 8.7)
# titles_1(document, subtitles[titles[7]][7], 8.8)

# Llamada a la función de formato de página
page_format(document, top_marg, left_marg, right_marg, bottom_marg, pag_height, pag_width, head_dist, foot_dist)

# Guardado del documento
document.save(r'C:\Users\user\Desktop\Dann Uc\Sistema de gestión de documentos\Pruebas\Instrumentos_Gestion_Final.docx')
