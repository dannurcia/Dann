"""Programa principal para la creación del documento Word Resumen Ejecutivo de un Proyecto
@Dann_Uc - 2022
"""


"""-------------------------------------------------- Módulos -------------------------------------------------------"""

from docx import Document
from docx.enum.section import WD_SECTION
from functions import *
from docx.enum.text import WD_ALIGN_PARAGRAPH


"""-------------------------------------- Variables de configuración de página --------------------------------------"""

top_marg = 2.5
left_marg = 3
right_marg = 3
bottom_marg = 2.5
pag_height = 297
pag_width = 210
head_dist = 1.5
foot_dist = 1


"""------------------------------------- Variables de configuración de párrafos -------------------------------------"""

i = 0                                           # Identacion izquierda del párrafo
fli = 0.5                                       # Identacion primera linea (sangria francesa para valores negativos)
sb = 0                                          # Espaciado antes del párrafo
sa = 6                                          # Espaciado antes del párrafo
ls = 1.2                                        # Interlineado
a = 3                                           # Alineación - 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
fn = 'Arial'                                    # Establece el tipo de letra
fs = 8                                          # Establece el tamaño de letra
name_style = 'Indent'                           # Nombre del estilo aplicado


"""-------------------------------------------- Variables del documento ---------------------------------------------"""

nombre_proyecto = 'Desarrollo de una herramienta computacional para la evaluación de los estados nutricionales ' \
                  'e hídrico de plantaciones de Palta Hass, a partir de las manifestaciones ópticas y térmicas ' \
                  'del cultivo, utilizando algoritmos de procesamiento de imágenes aéreas multiespectrales e ' \
                  'inteligencia artificial'
titles = [ 'Descripción', 'Objetivo', 'Datos Generales', 'Equipo Técnico']

subtitles = {titles[2]: ['Fecha de inicio', 'Fecha de cierre', 'Estado del proyecto'],
             titles[3]: ['Investigador principal', 'Co-Investigadores', 'Equipo Técnico de Apoyo',
                         'Co-Investigadores Asociados', 'Tesistas']}

descr = 'El presente proyecto apunta a desarrollar un aplicativo de software basado en algoritmos computacionales de ' \
        'procesamiento digital de imágenes aéreas adquiridas vía UAV, con el objetivo de estimar los estados ' \
        'nutricional e hídrico de plantaciones de paltas Hass. Esto permitirá a los productores tomar las acciones ' \
        'correctivas a tiempo con el propósito de proteger la plantación ante una posible enfermedad o plaga y ' \
        'mejorar la eficiencia en las dosificaciones de agua, fertilizantes y nutrientes para efectos de incrementar ' \
        'el rendimiento del cultivo y hacer un uso más racional de los recursos hídricos. La investigación se enfoca ' \
        'también en resolver el problema de monitoreo y estimación de los estados nutricionales e hídricos de las ' \
        'grandes plantaciones en el Perú, las cuales se han incrementado notablemente como consecuencia del gran ' \
        'crecimiento de las exportaciones de la persea americana Hass. En ese contexto, el proyecto involucra el ' \
        'estudio de las manifestaciones térmicas y ópticas de las deficiencias nutricionales e hídricas que puedan ' \
        'ser identificadas y caracterizadas a través de imágenes aéreas de grandes plantaciones, así como el ' \
        'desarrollo de técnicas computacionales de procesamiento de imágenes e inteligencia artificial. El proyecto ' \
        'está siendo desarrollado a través de la asociación entre el INICTEL-UNI y el Instituto Nacional de ' \
        'Innovación agraria (INIA), el cual brinda las áreas de plantaciones de monitoreo en la Estaciones de Donoso ' \
        'en Huaral y de La Molina en la ciudad de Lima. Asimismo ejecuta los procedimientos agronómicos pertinentes ' \
        'para el desarrollo de las actividades experimentales del proyecto.'

obj = 'Desarrollar un aplicativo de software basado en algoritmos computacionales de procesamiento digital de ' \
      'imágenes aéreas adquiridas vía UAV, con el objetivo de estimar los estados nutricional e hídrico de ' \
      'plantaciones de paltas Hass.'

fecha_inicio = 'Diciembre 2018'
fecha_cierre = 'Mayo 2021'
estado_proy = 'EN EJECUCIÓN'
invest_princ = 'Guillermo Kemper'
co_invest1 = 'Samuel Huamán'
co_invest2 = 'Grovher Palomino'
co_invest3 = 'Joel Telles'
co_invest4 = 'Daniel Arteaga'
co_invest5 = 'Itamar Salazar'
co_invest6 = 'Marco Apolinario'
equi_apoyo = 'Ivan Ortega (INICTEL-UNI)'
co_asocia1 = 'María Rojas (INIA)'
co_asocia2 = 'Rafael Calderón (INIA)'
co_asocia3 = 'Christian Del Carpio (FIM-UNI)'
co_asocia4 = 'Jonell Soto (INIA)'
co_asocia5 = 'José Oliden (FIM-UNI)'
tesista1 = 'Bruno Rivadeneyra (Ingeniería Nacional de Ingeniería)'
tesista2 = 'Lincol Vargas (Universidad Nacional de Ingeniería)'
tesista3 = 'Cesar Muñoz (Universidad Nacional de Ingeniería)'

paragraphs = [descr, obj, fecha_inicio, fecha_cierre, estado_proy, invest_princ, co_invest1, co_invest2, co_invest3,
              co_invest4, co_invest5, co_invest6, equi_apoyo, co_asocia1, co_asocia2, co_asocia3, co_asocia4,
              co_asocia5, tesista1, tesista2, tesista3]


"""------------------------------------------------- Programa -------------------------------------------------------"""

# Creación del documento
document = Document()

# Introduciendo el logo del encabezado
header = document.sections[0].header
paragraph = header.paragraphs[0]
logo_run = paragraph.add_run()
logo_run.add_picture("cabecera.png", width=Cm(15))

# Inicio de numeración de página
add_page_number(document.sections[0])

# Llamada a la función de estilo de párrafo
paragraph_re(document, i, fli, sb, sa, ls, a, fn, fs, name_style)

# Cuerpo del documento (títulos, subtítulos y párrafos)
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)                                      # Espaciado 12 pts antes del párrafo
p.paragraph_format.space_after = Pt(20)
p.paragraph_format.line_spacing = 1.5
p1 = p.add_run('Resumen Ejecutivo del Proyecto: "' + nombre_proyecto + '"')
p1.bold, p1.font.size, p1.font.name = True, Pt(10), 'Arial'

titles_0(document, titles[0], 1)
p1 = document.add_paragraph(paragraphs[0])
p1.style = document.styles[name_style]                                        # Aplicando el estilo al párrafo

titles_0(document, titles[1], 2)
p2 = document.add_paragraph(paragraphs[1])
p2.style = document.styles[name_style]                                        # Aplicando el estilo al párrafo
titles_0(document, titles[2], 3)
titles_1(document, subtitles[titles[2]][0], 3.1)
bullets_re(document, paragraphs[2])
titles_1(document, subtitles[titles[2]][1], 3.2)
bullets_re(document, paragraphs[3])
titles_1(document, subtitles[titles[2]][2], 3.3)
bullets_re(document, paragraphs[4])

titles_0(document, titles[3], 4)
titles_1(document, subtitles[titles[3]][0], 4.1)
bullets_re(document, paragraphs[5])
titles_1(document, subtitles[titles[3]][1], 4.2)
bullets_re(document, paragraphs[6])
bullets_re(document, paragraphs[7])
bullets_re(document, paragraphs[8])
bullets_re(document, paragraphs[9])
bullets_re(document, paragraphs[10])
bullets_re(document, paragraphs[11])
titles_1(document, subtitles[titles[3]][2], 4.3)
bullets_re(document, paragraphs[12])
titles_1(document, subtitles[titles[3]][3], 4.4)
bullets_re(document, paragraphs[13])
bullets_re(document, paragraphs[14])
bullets_re(document, paragraphs[15])
bullets_re(document, paragraphs[16])
bullets_re(document, paragraphs[17])
titles_1(document, subtitles[titles[3]][4], 4.5)
bullets_re(document, paragraphs[18])
bullets_re(document, paragraphs[19])
bullets_re(document, paragraphs[20])

# Llamada a la función de formato de página
page_format(document, top_marg, left_marg, right_marg, bottom_marg, pag_height, pag_width, head_dist, foot_dist)

# Guardado del documento
document.save(r'C:\Users\user\Desktop\Dann Uc\Sistema de gestión de documentos\Pruebas\Resumen_Ejecutivo.docx')



