"""Implementación de funciones para la creación de un documento Word con formato GdR
@Dann_Uc - 2022
"""


"""-------------------------------------------------- Módulos -------------------------------------------------------"""

from docx.shared import Inches, Pt, Cm, Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, Cm, Mm
from docx.oxml.ns import qn
from docx.shared import RGBColor
from datetime import date
from datetime import datetime
import re


"""---------------------------------------------- Funciones GdR------------------------------------------------------"""


def first_page(document, nEvi_form, dirOfAbrev, name_initials, nameEv, nameJef, date, cargo, year, dni, records_ind, descripcion_evi, full_name):
    """
    Diseño de la primera hoja del documento GdR
    :param document:
    :param nGdR:
    :param dirOfAbrev:
    :param nameColabAbrev:
    :param nameEv:
    :param nameJef:
    :param date:
    :param cargo:
    :param nameInd:
    :param nameColab:
    :param dni:
    :return:
    """
    document.add_picture('C:\\Users\\user\\PycharmProjects\\Dann\\cabecera.png',
                         width=Cm(15))  # Añadir imagen de cabecera
    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = p.add_run('INFORME GESTION DEL RENDIMIENTO N° ')
    p1.bold, p1.underline = True, True
    p1.font.size, p1.font.name = Pt(12), 'Arial'
    p2 = p.add_run(nEvi_form + ' - ' + year + ' - ')
    p2.italic, p2.bold, p2.underline = True, True, True
    p2.font.size, p2.font.name = Pt(12), 'Arial'
    p3 = p.add_run('INICTEL-UNI-')
    p3.bold, p3.underline = True, True
    p3.font.size, p3.font.name = Pt(12), 'Arial'
    p4 = p.add_run(dirOfAbrev + '-' + name_initials)
    p4.italic, p4.bold, p4.underline = True, True, True
    p4.font.size, p4.font.name = Pt(12), 'Arial'

    table = document.add_table(rows=1, cols=3)  # Creacion de una tabla de 1 fila y 3 columnas

    hdr_cells = table.rows[0].cells
    hdr00 = hdr_cells[0].paragraphs[0].add_run('A')
    hdr00.bold, hdr00.font.name, hdr00.font.size = False, 'Arial', Pt(10)
    hdr10 = hdr_cells[1].paragraphs[0].add_run(':')
    hdr10.bold, hdr10.font.name, hdr10.font.size = False, 'Arial', Pt(10)
    hdr20 = hdr_cells[2].paragraphs[0].add_run(nameJef)
    hdr20.bold, hdr20.italic, hdr20.font.name, hdr20.font.size = True, True, 'Arial', Pt(10)

    row_cells = table.add_row().cells  # Agregando una fila a la tabla
    r1c00 = row_cells[0].paragraphs[0].add_run('')
    r1c10 = row_cells[1].paragraphs[0].add_run('')
    r1c20 = row_cells[2].paragraphs[0].add_run(cargo)
    r1c20.bold, r1c20.italic, r1c20.font.name, r1c20.font.size = True, True, 'Arial', Pt(10)

    row_cells = table.add_row().cells  # Agregando una fila a la tabla
    r2c00 = row_cells[0].paragraphs[0].add_run('ASUNTO')
    r2c00.bold, r2c00.font.name, r2c00.font.size = False, 'Arial', Pt(10)
    r2c10 = row_cells[1].paragraphs[0].add_run(':')
    r2c10.bold, r2c10.font.name, r2c10.font.size = False, 'Arial', Pt(10)
    r2c20 = row_cells[2].paragraphs[0].add_run('Entregable PAO ' + year + ' - ' + nEvi_form)
    r2c20.bold, r2c20.italic, r2c20.font.name, r2c20.font.size = False, True, 'Arial', Pt(10)

    row_cells = table.add_row().cells  # Agregando una fila a la tabla
    r3c00 = row_cells[0].paragraphs[0].add_run('FECHA')
    r3c00.bold, r3c00.font.name, r3c00.font.size = False, 'Arial', Pt(10)
    r3c10 = row_cells[1].paragraphs[0].add_run(':')
    r3c10.bold, r3c10.font.name, r3c10.font.size = False, 'Arial', Pt(10)
    r3c20 = row_cells[2].paragraphs[0].add_run(date)
    r3c20.bold, r3c20.italic, r3c20.font.name, r3c20.font.size = False, True, 'Arial', Pt(10)

    # Configurando el ancho de las celdas para cada columna
    for b in list(range(3)):
        if b == 0:
            for a in list(range(4)):
                table.cell(a, b).width = Cm(2.5)
        elif b == 1:
            for a in list(range(4)):
                table.cell(a, b).width = Cm(1)
        elif b == 2:
            for a in list(range(4)):
                table.cell(a, b).width = Cm(11.5)

    q = document.add_paragraph(
        '_____________________________________________________________________________________________________')
    q.paragraph_format.alignment = 1

    r = document.add_paragraph()
    r.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r.paragraph_format.left_indent = Cm(0)
    r.paragraph_format.first_line_indent = Cm(2.5)
    r.paragraph_format.space_before = Pt(6)
    r.paragraph_format.space_after = Pt(6)
    r.paragraph_format.line_spacing = 1.5
    run1 = r.add_run('Me dirijo a usted para remitirle adjunto al presente el informe correspondiente '
                     'al indicador ')
    run1.font.size = Pt(11)
    run1.font.name = 'Arial'
    run1.bold = False
    run_nameInd = r.add_run('"' + str(records_ind) + '"')
    run_nameInd.font.size = Pt(11)
    run_nameInd.font.name = 'Arial'
    run_nameInd.bold = False
    run_nameInd.italic = True
    run2 = r.add_run(', denominada según detalle:')
    run2.font.size = Pt(11)
    run2.font.name = 'Arial'
    run2.bold = False

    s = document.add_paragraph()
    s.paragraph_format.left_indent = Cm(0.63)
    s.paragraph_format.first_line_indent = Cm(0.63)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(12)
    s.paragraph_format.alignment = 3
    s.paragraph_format.line_spacing = 1.5
    s.add_run('•  ')
    run3 = s.add_run('Un (01) ')
    run3.font.size = Pt(11)
    run3.font.name = 'Arial'
    run3.bold = False
    run4 = s.add_run(descripcion_evi.strip())
    run4.font.size = Pt(11)
    run4.font.name = 'Arial'
    run4.italic = True
    run4.bold = False
    run4 = s.add_run('. Actividad desarrollada hasta el presente mes del ' + year + '.')
    run4.font.size = Pt(11)
    run4.font.name = 'Arial'
    run4.bold = False

    t = document.add_paragraph()
    t.paragraph_format.left_indent = Cm(2.5)
    t.paragraph_format.first_line_indent = Cm(0.63)
    t.paragraph_format.space_before = Pt(12)
    t.paragraph_format.space_after = Pt(12)
    t.paragraph_format.alignment = 3
    t.paragraph_format.line_spacing = 1.5
    run5 = t.add_run('Sin otro particular quedo de usted.')
    run5.font.size = Pt(11)
    run5.font.name = 'Arial'
    run5.bold = False

    u = document.add_paragraph()
    u.paragraph_format.left_indent = Cm(2.5)
    u.paragraph_format.first_line_indent = Cm(0.63)
    u.paragraph_format.space_before = Pt(48)
    u.paragraph_format.space_after = Pt(48)
    u.paragraph_format.alignment = 3
    u.paragraph_format.line_spacing = 1.5
    run6 = u.add_run('Atentamente,')
    run6.font.size = Pt(11)
    run6.font.name = 'Arial'
    run6.bold = False

    table1 = document.add_table(rows=1, cols=2)  # Creacion de una tabla de 1 fila y 3 columnas
    table1.autofit = False
    hdr_cells = table1.rows[0].cells
    hdr00 = hdr_cells[0].paragraphs[0].add_run('')
    hdr10 = hdr_cells[1].paragraphs[0].add_run('_________________________________________________')

    row_cells = table1.add_row().cells  # Agregando una fila a la tabla
    r1c00 = row_cells[0].paragraphs[0].add_run('')
    r1c10 = row_cells[1].paragraphs[0].add_run(full_name)
    r1c10.bold, r1c10.italic, r1c10.font.name, r1c10.font.size = False, True, 'Arial', Pt(11)

    row_cells = table1.add_row().cells  # Agregando una fila a la tabla
    r1c00 = row_cells[0].paragraphs[0].add_run('')
    r1c10 = row_cells[1].paragraphs[0].add_run('DNI: ' + dni)
    r1c10.bold, r1c10.italic, r1c10.font.name, r1c10.font.size = False, True, 'Arial', Pt(11)

    table1.rows[0].height = Cm(0.1)
    for cell in table1.row_cells(0):
        cell.height = Cm(0.1)


def second_page(document, descripcion_evi, full_name, mes_actual, year):
    """
    Configuracion de la 2da página del documento
    :return:
    """
    document.add_picture('cabecera.png', width=Cm(15))  # Añadir imagen de cabecera
    a = document.add_paragraph('')
    a.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    a0 = a.add_run('DIRECCIÓN DE INVESTIGACIÓN Y DESARROLLO TECNOLÓGICO \n')
    a0.font.name = 'Cambria'
    a0.font.size = Pt(14)
    a0.font.bold = True
    a0.font.italic = True
    a1 = a.add_run('\n-  DIDT -')
    a1.font.size = Pt(14)
    a1.font.bold = True
    a1.font.italic = True
    a2 = a.add_run('\n \n \n')

    b = document.add_paragraph('\n')
    b.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    b1 = b.add_run('\n' + descripcion_evi.upper())
    b1.font.name = 'Cambria'
    b1.font.size = Pt(20)
    b1.font.bold = True
    b1.font.italic = True
    b2 = b.add_run('\n')

    c = document.add_paragraph('')
    c.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    c1 = c.add_run('Informe de Actividades')
    c1.font.name = 'Cambria'
    c1.font.size = Pt(16)
    c1.font.bold = True
    c1.font.italic = True
    c2 = c.add_run('\n \n \n')

    d = document.add_paragraph('')
    d.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    d1 = d.add_run('Autor:')
    d1.font.name = 'Cambria'
    d1.font.size = Pt(14)
    d1.font.bold = False
    d1.font.italic = False
    d2 = d.add_run('\n' + full_name)
    d2.font.name = 'Cambria'
    d2.font.size = Pt(14)
    d2.font.bold = False
    d2.font.italic = True

    e = document.add_paragraph('\n')
    e.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    e1 = e.add_run('\n SAN BORJA')
    e1.font.name = 'Cambria'
    e1.font.size = Pt(14)
    e1.font.bold = True
    e1.font.italic = True
    e2 = e.add_run('\n \n')
    e3 = e.add_run(mes_actual.upper() + ' – ' + year)
    e3.font.name = 'Cambria'
    e3.font.size = Pt(14)
    e3.font.bold = True
    e3.font.italic = True


def page_format(document, top_marg, left_marg, right_marg, bottom_marg, pag_height, pag_width, head_dist, foot_dist):
    """
    Configuración de la página en formato A4
    :param document:
    :param top_marg:
    :param left_marg:
    :param right_marg:
    :param bottom_marg:
    :param pag_height:
    :param pag_width:
    :param head_dist:
    :param foot_dist:
    :return:
    """
    for sec in document.sections:
        # Márgenes de página (25 mm arriba a la izquierda y 15 mm abajo a la derecha)
        sec.top_margin = Cm(top_marg)
        sec.left_margin = Cm(left_marg)
        sec.right_margin = Cm(right_marg)
        sec.bottom_margin = Cm(bottom_marg)
        # Configura el papel al formato A4
        sec.page_height = Mm(pag_height)
        sec.page_width = Mm(pag_width)
        # Establece la distancia del pie de página del encabezado
        sec.header_distance = Cm(head_dist)
        sec.footer_distance = Cm(foot_dist)


def title_0(document, title, num):
    """
    Configuración de los títulos del primer nivel
    :param document:
    :param title:
    :param num:
    :return:
    """

    heading = document.add_heading(str(num) + '.  ' + title)
    title_style = heading.style
    rFonts = title_style.element.rPr.rFonts                 # lxml de la fuente
    rFonts.set(qn("w:asciiTheme"), "Times New Roman")       # Configura fuente xml
    title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title_style.paragraph_format.left_indent = Cm(0.63)
    # title_style = document.add_paragraph(str(num) + '.  ' + title)
    # title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # title_style.paragraph_format.line_spacing = 1.2
    # title_style.paragraph_format.left_indent = Cm(0.5)
    title_style.paragraph_format.space_before = Pt(12)
    title_style.paragraph_format.space_after = Pt(12)
    # title_style.paragraph_format.first_line_indent = Cm(0.63)
    font = title_style.font                                     # Estilo de fuente del titulo
    font.name = 'Times New Roman'                               # Necesario, acorde al xml
    font.size = Pt(14)                                          # Tamaño de fuente
    font.bold = True

    return 'title_0', str(num) + '.  ' + title


def title_1(document, title, num):
    """
    Configuración de los títulos del segundo nivel
    :param document:
    :param title:
    :param num:
    :return:
    """
    heading = document.add_heading(str(num) + '.  ' + title, level=2)
    title_style = heading.style
    rFonts = title_style.element.rPr.rFonts                     # lxml de la fuente
    rFonts.set(qn("w:asciiTheme"), "Times New Roman")           # Configura fuente xml
    title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title_style.paragraph_format.left_indent = Cm(0.63)
    # title_style = document.add_paragraph(str(num) + '.  ' + title)
    # title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # title_style.paragraph_format.line_spacing = 1.2
    # title_style.paragraph_format.left_indent = Cm(0.5)
    title_style.paragraph_format.space_before = Pt(6)
    title_style.paragraph_format.space_after = Pt(6)
    # title_style.paragraph_format.first_line_indent = Cm(0.63)
    font = title_style.font                                     # Estilo de fuente del titulo
    font.name = 'Cambria'                                       # Necesario, acorde al xml
    font.size = Pt(12)                                          # Tamaño de fuente
    font.bold = True

    return 'title_1', str(num) + '.  ' + title


def title_2(document, title, num, sub):
    """
    Configuración de los títulos del tercer nivel
    :param document:
    :param title:
    :param num:
    :return:
    """
    heading = document.add_heading(str(num) + '.' + str(sub) + '.  ' + title, level=3)
    title_style = heading.style
    rFonts = title_style.element.rPr.rFonts                     # lxml de la fuente
    rFonts.set(qn("w:asciiTheme"), "Times New Roman")           # Configura fuente xml
    title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title_style.paragraph_format.left_indent = Cm(0.63)
    # title_style = document.add_paragraph(str(num) + '.  ' + title)
    # title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # title_style.paragraph_format.line_spacing = 1.2
    # title_style.paragraph_format.left_indent = Cm(0.5)
    title_style.paragraph_format.space_before = Pt(6)
    title_style.paragraph_format.space_after = Pt(6)
    # title_style.paragraph_format.first_line_indent = Cm(0.63)
    font = title_style.font                                     # Estilo de fuente del titulo
    font.name = 'Cambria'                               # Necesario, acorde al xml
    font.size = Pt(12)                                          # Tamaño de fuente
    font.bold = False
    font.italic = False

    return 'title_2', str(num) + '.  ' + title


def paragraph_0(document, i, fli, sb, sa, ls, a, fn, fs, name_style):
    """
    Formato para los párrafos del documento GdR
    :param document:
    :param i:
    :param fli:
    :param sb:
    :param sa:
    :param ls:
    :param a:
    :param fn:
    :param fs:
    :param name_style:
    :return:
    """
    style = document.styles.add_style(name_style, WD_STYLE_TYPE.PARAGRAPH)      # Acceso al diccionario de estilos
    paragraph_format = style.paragraph_format                                   # Propiedad de ParagraphFormat
    paragraph_format.left_indent = Inches(i)                                    # Identacion izquierda del párrafo
    paragraph_format.first_line_indent = Cm(fli)                                # Identacion primera linea (sangria francesa)
    paragraph_format.space_before = Pt(sb)                                      # Espaciado 12 pts antes del párrafo
    paragraph_format.space_after = Pt(sa)                                       # Espaciado 12 pts después del párrafo
    paragraph_format.widow_control = True                                       # Control de viudas y huérfanos
    paragraph_format.line_spacing = ls                                          # Interlineado a 1.5 líneas
    paragraph_format.alignment = a                                              # 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
    font = style.font                                                           # Definiendo el formato del caracter
    font.name = fn                                                              # Establece el tipo de letra
    font.size = Pt(fs)                                                          # Establece el tamaño de letra


def items(document, text):
    """
    Configuración de los items del primer nivel
    :param document:
    :param text:
    :return:
    """
    t2 = document.add_paragraph()
    run = t2.add_run('-  ' + text)
    # run = t2.add_run('●  ' + text)
    # run = t2.add_run('➤  ' + text)
    # run = t2.add_run('➢  ' + text)
    run.font.size = Pt(12)
    run.font.name = 'Calibri(Cuerpo)'
    t2.paragraph_format.left_indent = Cm(0.63)
    t2.paragraph_format.first_line_indent = Cm(0.63)
    t2.paragraph_format.space_before = Pt(0)
    t2.paragraph_format.space_after = Pt(6)
    t2.paragraph_format.alignment = 3
    t2.paragraph_format.line_spacing = 1
    run.bold = False


def bullets(document, text):
    """
    Configuración de las viñetas en el segundo nivel
    :param document:
    :param text:
    :return:
    """
    t3 = document.add_paragraph()
    # run = t2.add_run('-  ' + text)
    run = t3.add_run('•  ' + text)
    # run = t2.add_run('➤  ' + text)
    # run = t2.add_run('➢  ' + text)
    run.font.size = Pt(12)
    run.font.name = 'Calibri(Cuerpo)'
    t3.paragraph_format.left_indent = Cm(1.9)
    t3.paragraph_format.first_line_indent = Cm(0.63)
    t3.paragraph_format.space_before = Pt(0)
    t3.paragraph_format.space_after = Pt(6)
    t3.paragraph_format.alignment = 3
    t3.paragraph_format.line_spacing = 1
    run.bold = False


def ref_bullets(document, text, num_ref):
    """
    Configuración de las viñetas en el segundo nivel
    :param document:
    :param text:
    :return:
    """
    t3 = document.add_paragraph()
    run = t3.add_run('[' + str(num_ref) + ']  ' + text)
    run.font.size = Pt(12)
    run.font.name = 'Calibri(Cuerpo)'
    t3.paragraph_format.left_indent = Cm(1.9)
    t3.paragraph_format.first_line_indent = Cm(-1.0)
    t3.paragraph_format.space_before = Pt(0)
    t3.paragraph_format.space_after = Pt(6)
    t3.paragraph_format.alignment = 3
    t3.paragraph_format.line_spacing = 1
    run.bold = False


# Tablas e imágenes

def get_tabla(docum, num_tabla, tit_tabla, lista_elem):
    # PONIENDO EL TITULO
    parrafo = docum.add_paragraph('Tabla ' + str(num_tabla) + ". " + tit_tabla)
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.paragraph_format.space_before = Pt(12)
    parrafo.paragraph_format.space_after = Pt(6)
    ru = parrafo.runs[0]
    ru.font.size = Pt(12)
    ru.font.name = 'Calibri'

    vvv = lista_elem.split("||")
    # definiendo la dimension de la tabla
    num_rows = len(vvv)
    num_cols = len(vvv[0].split("$$"))
    tablita = docum.add_table(num_rows, num_cols)
    tablita.autofit = True
    tablita.alignment = WD_TABLE_ALIGNMENT.CENTER
    cont = 0
    for k in vvv:
        fff = k.split("$$")
        rrr = tablita.rows[cont].cells
        n_o = 0
        for h in fff:

            rrr[n_o].text = h

            ppp = rrr[n_o].paragraphs[0]
            ppp.paragraph_format.space_before = Pt(0)
            ppp.paragraph_format.space_after = Pt(6)
            mmm = ppp.runs[0]
            mmm.font.size = Pt(12)
            mmm.font.name = 'Calibri'

            if (cont == 0):
                ppp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            n_o = n_o + 1

        cont = cont + 1

    tablita.style = 'Table Grid'

    # agregando un espaciado despues de la tabla
    p_esp = docum.add_paragraph()
    p_esp.paragraph_format.space_before = Pt(12)
    p_esp.paragraph_format.space_after = Pt(0)
    r_esp = p_esp.add_run('')
    r_esp.font.size = Pt(12)
    r_esp.font.name = 'Calibri'


def get_imagen(docum, path_img, ancho_im, num_im, tit_im):
    #agregando un espaciado antes de la imagen
    p_esp = docum.add_paragraph()
    p_esp.paragraph_format.space_before = Pt(12)
    p_esp.paragraph_format.space_after = Pt(0)
    r_esp = p_esp.add_run('')
    r_esp.font.size = Pt(12)
    r_esp.font.name = 'Calibri'

    docum.add_picture(path_img, width = Cm(ancho_im))
    last_paragraph = docum.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Poniendo el titulo
    parrafo = docum.add_paragraph()
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.paragraph_format.space_before = Pt(0)
    parrafo.paragraph_format.space_after = Pt(6)
    r = parrafo.add_run('Figura ' + str(num_im) + ". " + tit_im)
    r.font.size = Pt(12)
    r.font.name = 'Calibri'


# def initials():
#     """
#     Retorna las iniciales de los nombres del usuario
#     :return:
#     """
#     pattern = r'[A-Z]+'
#     coincid = re.findall()


# Funciones del pie de página

def create_element(name):
    """
    Crea un elemento xml
    :param name:
    :return:
    """
    return OxmlElement(name)


def create_attribute(element, name, value):
    """
    Atributos del elemento xml
    :param element:
    :param name:
    :param value:
    :return:
    """
    element.set(ns.qn(name), value)


def add_page_number(doc_sec):
    """
    Agrega la numeración del pie de página a partir de la segunda sección creada en el documento(seccion1)
    :param doc_sec:
    :return:
    """
    doc_sec.footer.is_linked_to_previous = False
    doc_sec.footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT                # Alineacion del pie de página

    page_num_run = doc_sec.footer.paragraphs[0].add_run()
    page_num_run.font.size = Pt(12)
    page_num_run.font.name = 'Calibri'

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)


def mes2words(mes):
    """

    :param mes: tipo str
    :return:
    """
    meses = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio',
             'Agosto', 'Setiembre', 'Octubre', 'Noviembre', 'Diciembre']

    if mes=='01':
        mes_word = meses[0]
    elif mes=='02':
        mes_word = meses[1]
    elif mes=='03':
        mes_word = meses[2]
    elif mes=='04':
        mes_word = meses[3]
    elif mes=='05':
        mes_word = meses[4]
    elif mes=='06':
        mes_word = meses[5]
    elif mes=='07':
        mes_word = meses[6]
    elif mes=='08':
        mes_word = meses[7]
    elif mes=='09':
        mes_word = meses[8]
    elif mes=='10':
        mes_word = meses[9]
    elif mes=='11':
        mes_word = meses[10]
    elif mes=='12':
        mes_word = meses[11]

    return mes_word


"""---------------------------------- Funciones Instrumentos de Gestion Final ---------------------------------------"""


def first_page_re(document):
    """

    :param document:
    :return:
    """
    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n \n \n \n')
    p1 = p.add_run('Documento Final')
    p1.bold, p1.underline = False, False
    p1.font.size, p1.font.name = Pt(18), 'Arial'
    p.add_run('\n \n')
    p2 = p.add_run('PRIORIDAD ANUAL DE GESTIÓN DEL ÓRGANO O UNIDAD ORGÁNICA N°4')
    p2.bold, p2.underline = False, False
    p2.font.size, p2.font.name = Pt(18), 'Arial'
    p.add_run('\n \n')

    heading = document.add_heading('GUIA METODOLOGICA PARA LA FORMULACIÓN DE PROYECTOS DE INVESTIGACIÓN Y '
                                   'DESARROLLO TECNOLOGICO EN DIDT', level=2)
    title_style = heading.style
    rFonts = title_style.element.rPr.rFonts  # lxml de la fuente
    rFonts.set(qn("w:asciiTheme"), "Arial")  # Configura fuente xml
    title_style.font.color.rgb = RGBColor(0x00, 0x70, 0xc0)
    title_style.paragraph_format.left_indent = Cm(0.63)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_before = Pt(6)
    title_style.paragraph_format.space_after = Pt(6)
    font = title_style.font  # Estilo de fuente del titulo
    font.name = 'Arial'  # Necesario, acorde al xml
    font.size = Pt(14)  # Tamaño de fuente
    font.bold = True

    q = document.add_paragraph()
    paragraph_format = q.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q.add_run('\n \n \n \n \n \n')
    # today = date.today()
    # fecha_doc = '{}'.format(today.day) + ' de ' + '{}'.format(today.month) + ' {}'.format(today.year)
    months = ("Enero", "Febrero", "Marzo", "Abri", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre",
              "Noviembre", "Diciembre")
    date1 = date.today()
    day = date1.day
    month = date1.month
    year = date1.year
    message = "{} de {} del {}".format(day, month, year)
    q1 = q.add_run(str(message))
    q1.bold, q1.underline = False, False
    q1.font.size, q1.font.name = Pt(18), 'Arial'
    q.add_run('\n \n \n \n \n \n')
    q2 = q.add_run('Ing. Joel Telles Castillo\n')
    q2.bold, q2.underline = False, False
    q2.font.size, q2.font.name = Pt(10), 'Arial'
    q3 = q.add_run('Director de Investigación y Desarrollo Tecnologico INICTEL-UNI')
    q3.bold, q3.underline = False, False
    q3.font.size, q3.font.name = Pt(10), 'Arial'


def titles_0(document, title, num):
    """
    Configuración de los títulos del primer nivel
    :param document:
    :param title:
    :param num:
    :return:
    """
    heading = document.add_heading(str(num) + '.   ' + title)
    title_style = heading.style
    rFonts = title_style.element.rPr.rFonts                 # lxml de la fuente
    rFonts.set(qn("w:asciiTheme"), "Arial")       # Configura fuente xml
    title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title_style.paragraph_format.left_indent = Cm(0)
    # title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_style.paragraph_format.line_spacing = 1.5
    title_style.paragraph_format.space_before = Pt(12)
    title_style.paragraph_format.space_after = Pt(0)
    # title_style.paragraph_format.first_line_indent = Cm(0.63)
    font = title_style.font                                     # Estilo de fuente del titulo
    font.name = 'Arial'                               # Necesario, acorde al xml
    font.size = Pt(10)                                          # Tamaño de fuente
    font.bold = True

    return 'title_0', str(num) + '.  ' + title


def titles_1(document, text, num):
    """
    Configuración de los títulos del segundo nivel
    :param document:
    :param title:
    :param num:
    :return:
    """
    t2 = document.add_paragraph()
    run = t2.add_run(str(num) + '.  ' + text)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    t2.paragraph_format.left_indent = Cm(0.5)
    t2.paragraph_format.first_line_indent = Cm(0.63)
    t2.paragraph_format.space_before = Pt(3)
    t2.paragraph_format.space_after = Pt(6)
    t2.paragraph_format.alignment = 3
    t2.paragraph_format.line_spacing = 1
    run.bold = True


def titles_2(document, text, num):
    """
    Configuración de los títulos del segundo nivel
    :param document:
    :param title:
    :param num:
    :return:
    """
    t2 = document.add_paragraph()
    run = t2.add_run(str(num) + '.  ' + text)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    t2.paragraph_format.left_indent = Cm(0.63)
    t2.paragraph_format.first_line_indent = Cm(0.8)
    t2.paragraph_format.space_before = Pt(0)
    t2.paragraph_format.space_after = Pt(0)
    t2.paragraph_format.alignment = 3
    t2.paragraph_format.line_spacing = 1
    run.bold = False

def titles_3(document, text, num, sub):
    """
    Configuración de los títulos del segundo nivel
    :param document:
    :param title:
    :param num:
    :return:
    """
    t2 = document.add_paragraph()
    run = t2.add_run(str(num) + '.' + str(sub) + '.  ' + text)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    t2.paragraph_format.left_indent = Cm(0.63)
    t2.paragraph_format.first_line_indent = Cm(0.8)
    t2.paragraph_format.space_before = Pt(0)
    t2.paragraph_format.space_after = Pt(0)
    t2.paragraph_format.alignment = 3
    t2.paragraph_format.line_spacing = 1
    run.bold = False


def paragraph_re(document, i, fli, sb, sa, ls, a, fn, fs, name_style):
    """
    Formato para los párrafos del documento Resumen Ejecutivo
    :param document:
    :param i:
    :param fli:
    :param sb:
    :param sa:
    :param ls:
    :param a:
    :param fn:
    :param fs:
    :param name_style:
    :return:
    """
    style = document.styles.add_style(name_style, WD_STYLE_TYPE.PARAGRAPH)      # Acceso al diccionario de estilos
    paragraph_format = style.paragraph_format                                   # Propiedad de ParagraphFormat
    paragraph_format.left_indent = Inches(i)                                    # Identacion izquierda del párrafo
    paragraph_format.first_line_indent = Cm(fli)                                # Identacion primera linea (sangria francesa)
    paragraph_format.space_before = Pt(sb)                                      # Espaciado 12 pts antes del párrafo
    paragraph_format.space_after = Pt(sa)                                       # Espaciado 12 pts después del párrafo
    paragraph_format.widow_control = True                                       # Control de viudas y huérfanos
    paragraph_format.line_spacing = ls                                          # Interlineado a 1.5 líneas
    paragraph_format.alignment = a                                              # 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
    font = style.font                                                           # Definiendo el formato del caracter
    font.name = fn                                                              # Establece el tipo de letra
    font.size = Pt(fs)                                                          # Establece el tamaño de letra


def bullets_re(document, text):
    """
    Configuración de las viñetas en el segundo nivel
    :param document:
    :param text:
    :return:
    """
    t3 = document.add_paragraph()
    # run = t2.add_run('-  ' + text)
    run = t3.add_run('•  ' + text)
    # run = t2.add_run('➤  ' + text)
    # run = t2.add_run('➢  ' + text)
    run.font.size = Pt(8)
    run.font.name = 'Arial'
    t3.paragraph_format.left_indent = Cm(1.75)
    t3.paragraph_format.first_line_indent = Cm(0.63)
    t3.paragraph_format.space_before = Pt(0)
    t3.paragraph_format.space_after = Pt(12)
    t3.paragraph_format.alignment = 3
    t3.paragraph_format.line_spacing = 1
    run.bold = False


"""----------------------------------------- Funciones Posprocesamiento ---------------------------------------------"""


def correcciones(lineas):
    """

    :param lineas:
    :return:
    """
    resultado = []

    for linea in lineas:
        # Reemplazar varios espacios entre palabras por uno solo
        linea = re.sub(r"(\S) {2,}", r"\1 ", linea)
        # Añadir espacio tras signo de puntuación no final
        linea = re.sub(r"([:;.,]+)(\S)", r"\1 \2", linea)
        # Quitar salto de línea
        linea = re.sub(r"([:;.,]+)(\S)", r"\1 \2", linea)

        resultado.append(linea)

    return resultado